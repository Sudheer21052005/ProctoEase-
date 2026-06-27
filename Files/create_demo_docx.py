from docx import Document
from docx.shared import Pt

output_path = r"c:\Users\Dell\OneDrive\Desktop\ProctoEase Mini Project\ProctoEase\Project_Demo_Viva_Script.docx"

title = "ProctoEase: Demo/Viva Presentation Script"

sections = {
    "1. Introduction": (
        "Good morning/afternoon. I am presenting ProctoEase, which is an online exam and recruitment platform. "
        "The main idea is to conduct technical assessments in a secure, fair, and scalable way. "
        "Candidates can take coding and theory exams online, and recruiters can monitor progress, review results, and make better hiring decisions from one system."
    ),
    "2. Core Features": (
        "Our backend has three important engines. "
        "First is Judge0 integration, which allows candidates to run code in a safe execution environment and get real outputs. "
        "Second is the proctoring system, which tracks suspicious behavior like tab switching and other exam violations. "
        "Third is the risk scoring system, which analyzes candidate activity and gives an overall risk level. "
        "On the frontend, we provide role-based dashboards for candidate, recruiter, and admin users so each user sees only what is relevant to them."
    ),
    "3. Improvements Through Phases": (
        "After building the backend, I improved the project in seven structured phases.\n\n"
        "In Phase 1, I fixed core reliability issues: candidates can now resume exams without duplicate attempts, code execution status and polling are accurate, and dashboard/reporting numbers are correct.\n\n"
        "In Phase 2, I cleaned the frontend architecture by following API to Hook to UI flow. I removed direct API calls from components, which made the code cleaner and easier to maintain.\n\n"
        "In Phase 3, I built a recruiter workspace with dedicated sections for Details, Questions, Attempts, Analytics, and Proctoring. This gave recruiters a single control center.\n\n"
        "In Phase 4, I added advanced capabilities: risk scoring UI, deeper proctoring insights, candidate answer review, and a manual grading draft flow for subjective evaluation.\n\n"
        "In Phase 5, I improved state management: Zustand handles client-side state, and React Query handles server data. This reduced confusion and data duplication.\n\n"
        "In Phase 6, I optimized performance using lazy loading, server-side pagination, virtualized lists, and improved cache settings to make the app faster and more responsive.\n\n"
        "In Phase 7, I focused on security and production readiness with stronger RBAC checks, backend authorization fixes for sensitive endpoints, frontend guard handling for forbidden actions, and a dedicated security QA script."
    ),
    "4. What Makes This Project Unique": (
        "What makes this project strong is that it combines online exams, secure code execution, proctoring intelligence, and risk-based analysis in one platform. "
        "It is not just an exam portal; it is a complete recruitment assessment workflow. "
        "Also, the project was improved phase by phase in a disciplined way, so it is not only feature-rich but also maintainable, performant, and secure."
    ),
    "5. Conclusion": (
        "To conclude, ProctoEase solves a real hiring problem by making online technical assessments more trustworthy and easier to manage. "
        "I started from core backend capabilities and then systematically improved correctness, architecture, recruiter workflow, advanced insights, state handling, performance, and security. "
        "Today, the system is much closer to a stable production-ready platform for modern recruitment use cases. Thank you."
    ),
}

doc = Document()
doc.add_heading(title, level=0)

for heading, body in sections.items():
    doc.add_heading(heading, level=1)
    for para in body.split("\n\n"):
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(10)

doc.save(output_path)
print(output_path)
